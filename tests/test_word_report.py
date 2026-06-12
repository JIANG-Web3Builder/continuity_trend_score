import pandas as pd
from docx import Document

from scripts.generate_word_report import collect_report_data, generate_report


REPORT_CLOSES = [100, 104, 108, 110, 109, 108, 107, 106, 105, 104, 103, 112, 118, 122, 120, 118, 116, 114, 112, 110, 109, 120, 128, 136]


def _write_symbol_csv(root, group, symbol, closes):
    group_dir = root / group
    group_dir.mkdir(parents=True, exist_ok=True)
    dates = pd.date_range("2024-01-01", periods=len(closes), freq="D")
    frame = pd.DataFrame(
        {
            "ts_code": symbol,
            "trade_date": dates.strftime("%Y%m%d"),
            "open": closes,
            "high": [value + 1 for value in closes],
            "low": [value - 1 for value in closes],
            "close": closes,
            "pre_close": [closes[0], *closes[:-1]],
            "change": [0.0, *[closes[index] - closes[index - 1] for index in range(1, len(closes))]],
            "pct_chg": 0.0,
            "vol": [1000 + index * 20 for index in range(len(closes))],
            "amount": [close * (1000 + index * 20) for index, close in enumerate(closes)],
        }
    )
    frame.to_csv(group_dir / f"{symbol}.csv", index=False)


def test_generate_word_report_creates_docx_and_chart_assets(tmp_path):
    _write_symbol_csv(tmp_path, "index", "000001.SH", REPORT_CLOSES)
    _write_symbol_csv(tmp_path, "sector", "BK1128.DC", [round(value * (1 + 0.01 * (index % 5)), 2) for index, value in enumerate(REPORT_CLOSES)])
    _write_symbol_csv(tmp_path, "commodity", "AUL.SHF", [round(value * (0.95 + 0.005 * index), 2) for index, value in enumerate(REPORT_CLOSES)])

    output_path = tmp_path / "report.docx"
    chart_dir = tmp_path / "assets"
    result = generate_report(data_dir=tmp_path, output_path=output_path, chart_dir=chart_dir)
    document = Document(output_path)
    paragraphs = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert result == output_path
    assert output_path.exists()
    assert len(list(chart_dir.glob("*.png"))) >= 5
    assert len(document.inline_shapes) >= 5
    assert len(document.tables) <= 4
    assert "研究背景与问题定义" in paragraphs
    assert "趋势评分的真实驱动因素" in paragraphs
    assert "指数样本分析" in paragraphs
    assert "板块样本分析" in paragraphs
    assert "商品样本分析" in paragraphs
    assert "指数样本分析\n覆盖" in paragraphs
    assert "板块样本分析\n覆盖" in paragraphs
    assert "商品样本分析\n覆盖" in paragraphs
    assert "附录" in paragraphs
    assert "鐮旂┒" not in paragraphs


def test_collect_report_data_contains_research_analysis_layers(tmp_path):
    _write_symbol_csv(tmp_path, "index", "000001.SH", REPORT_CLOSES)
    _write_symbol_csv(tmp_path, "sector", "BK1128.DC", [round(value * (1 + 0.01 * (index % 5)), 2) for index, value in enumerate(REPORT_CLOSES)])
    _write_symbol_csv(tmp_path, "commodity", "AUL.SHF", [round(value * (0.95 + 0.005 * index), 2) for index, value in enumerate(REPORT_CLOSES)])

    analysis = collect_report_data(tmp_path)

    assert {
        "total_return",
        "annual_volatility",
        "waves_per_year",
        "min_reversal_pct",
        "reversal_vol_multiple",
    }.issubset(analysis["assets"].columns)
    assert {"feature", "correlation"}.issubset(analysis["score_correlations"].columns)
    assert {"case_type", "symbol", "total_score"}.issubset(analysis["typical_waves"].columns)
    assert not analysis["wave_distribution"].empty

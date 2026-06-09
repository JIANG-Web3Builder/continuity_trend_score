from __future__ import annotations

import argparse
import math
import sys
from pathlib import Path
from typing import Any

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT_DIR = Path(__file__).resolve().parents[1]
SRC_DIR = ROOT_DIR / "src"
if str(SRC_DIR) not in sys.path:
    sys.path.insert(0, str(SRC_DIR))

from trend_score.catalog import ASSET_GROUPS, get_assets
from trend_score.compare import score_interval_continuity
from trend_score.data import load_symbol_data
from trend_score.scoring import score_waves
from trend_score.waves import _adaptive_reversal_threshold, detect_waves


REPORT_TITLE = "连续性系统—品种历史波段趋势识别与强度评分研究报告"
DEFAULT_OUTPUT = ROOT_DIR / "reports" / "continuity_wave_research_report.docx"
DEFAULT_CHART_DIR = ROOT_DIR / "reports" / "assets"
RECENT_WINDOW_DAYS = 120


def generate_report(
    data_dir: str | Path = ROOT_DIR / "data",
    output_path: str | Path = DEFAULT_OUTPUT,
    chart_dir: str | Path = DEFAULT_CHART_DIR,
) -> Path:
    data_dir = Path(data_dir)
    output_path = Path(output_path)
    chart_dir = Path(chart_dir)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    chart_dir.mkdir(parents=True, exist_ok=True)

    analysis = collect_report_data(data_dir)
    chart_paths = render_charts(analysis, chart_dir)
    write_docx_report(analysis, chart_paths, output_path)
    return output_path


def collect_report_data(data_dir: Path) -> dict[str, Any]:
    asset_rows: list[dict[str, Any]] = []
    wave_rows: list[dict[str, Any]] = []
    interval_rows: list[dict[str, Any]] = []
    price_frames: dict[str, pd.DataFrame] = {}
    group_frames: dict[str, list[pd.DataFrame]] = {group: [] for group in ASSET_GROUPS}

    for group in ASSET_GROUPS:
        for asset in get_assets(group):
            csv_path = data_dir / group / f"{asset.ts_code}.csv"
            if not csv_path.exists():
                continue

            prices = load_symbol_data(data_dir, asset.ts_code, group=group)
            if prices.empty:
                continue

            prices = prices.sort_values("trade_date").reset_index(drop=True)
            threshold = _adaptive_reversal_threshold(prices, 1.2)
            waves = detect_waves(prices, symbol=asset.ts_code)
            scored = score_waves(prices, waves)

            price_frames[asset.ts_code] = prices
            group_frames[group].append(prices)
            asset_rows.append(_asset_summary_row(group, asset.ts_code, asset.name, prices, scored, threshold))
            wave_rows.extend(_wave_distribution_rows(group, asset.ts_code, asset.name, scored))

    for group, frames in group_frames.items():
        if frames:
            interval_rows.extend(_interval_rows(group, pd.concat(frames, ignore_index=True)))

    assets = pd.DataFrame(asset_rows)
    wave_distribution = pd.DataFrame(wave_rows)
    intervals = pd.DataFrame(interval_rows)

    group_summary = _group_summary(assets)
    score_correlations = _score_correlations(wave_distribution)
    typical_waves = _select_typical_waves(wave_distribution)
    typical_paths = _typical_wave_paths(price_frames, typical_waves)

    top_waves = (
        wave_distribution.sort_values("total_score", ascending=False).head(20).reset_index(drop=True)
        if not wave_distribution.empty
        else wave_distribution
    )
    intervals = (
        intervals.sort_values("interval_score", ascending=False).reset_index(drop=True)
        if not intervals.empty
        else intervals
    )

    return {
        "data_dir": data_dir,
        "generated_at": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M"),
        "assets": assets,
        "group_summary": group_summary,
        "wave_distribution": wave_distribution,
        "score_correlations": score_correlations,
        "typical_waves": typical_waves,
        "typical_paths": typical_paths,
        "top_waves": top_waves,
        "intervals": intervals,
    }


def _asset_summary_row(
    group: str,
    symbol: str,
    name: str,
    prices: pd.DataFrame,
    scored: pd.DataFrame,
    threshold: float,
) -> dict[str, Any]:
    closes = prices["close"].astype(float)
    returns = closes.pct_change().dropna()
    daily_abs_move = closes.diff().abs().dropna()
    years = max(len(prices) / 252.0, 1 / 252.0)
    total_return = (closes.iloc[-1] / closes.iloc[0] - 1.0) * 100.0 if closes.iloc[0] else 0.0
    annual_volatility = float(returns.std() * math.sqrt(252) * 100.0) if len(returns) > 1 else 0.0
    median_close = float(closes.median()) if not closes.empty else 0.0
    median_abs_move = float(daily_abs_move.median()) if not daily_abs_move.empty else 0.0
    threshold_pct = float(threshold / median_close * 100.0) if median_close else 0.0
    vol_multiple = float(threshold / median_abs_move) if median_abs_move else 0.0

    return {
        "group": group,
        "group_name": ASSET_GROUPS[group],
        "symbol": symbol,
        "name": name,
        "rows": int(len(prices)),
        "start_date": prices["trade_date"].min().strftime("%Y-%m-%d"),
        "end_date": prices["trade_date"].max().strftime("%Y-%m-%d"),
        "min_reversal": round(float(threshold), 2),
        "min_reversal_pct": round(float(threshold_pct), 2),
        "reversal_vol_multiple": round(float(vol_multiple), 2),
        "waves": int(len(scored)),
        "up_waves": int((scored["direction"] == "up").sum()) if not scored.empty else 0,
        "down_waves": int((scored["direction"] == "down").sum()) if not scored.empty else 0,
        "avg_days": round(float(scored["days"].mean()), 1) if not scored.empty else 0.0,
        "median_points": round(float(scored["points"].median()), 2) if not scored.empty else 0.0,
        "max_points": round(float(scored["points"].max()), 2) if not scored.empty else 0.0,
        "top_score": round(float(scored["total_score"].max()), 2) if not scored.empty else 0.0,
        "total_return": round(float(total_return), 2),
        "annual_volatility": round(float(annual_volatility), 2),
        "trading_years": round(float(years), 2),
        "waves_per_year": round(float(len(scored) / years), 2),
    }


def _wave_distribution_rows(group: str, symbol: str, name: str, scored: pd.DataFrame) -> list[dict[str, Any]]:
    if scored.empty:
        return []

    rows: list[dict[str, Any]] = []
    for _, wave in scored.iterrows():
        pct_change = float(wave["pct_change"])
        days = max(int(wave["days"]), 1)
        rows.append(
            {
                "group": group,
                "group_name": ASSET_GROUPS[group],
                "symbol": symbol,
                "name": name,
                "direction": wave["direction"],
                "direction_name": "上涨" if wave["direction"] == "up" else "下跌",
                "level": wave["level"],
                "start_date": pd.Timestamp(wave["start_date"]).strftime("%Y-%m-%d"),
                "end_date": pd.Timestamp(wave["end_date"]).strftime("%Y-%m-%d"),
                "start_price": float(wave["start_price"]),
                "end_price": float(wave["end_price"]),
                "points": round(float(wave["points"]), 2),
                "pct_change": round(pct_change, 2),
                "abs_pct_change": round(abs(pct_change), 2),
                "days": days,
                "points_per_day": round(float(wave["points"]) / days, 4),
                "pct_per_day": round(abs(pct_change) / days, 4),
                "strength_score": round(float(wave["strength_score"]), 2),
                "duration_score": round(float(wave["duration_score"]), 2),
                "slope_score": round(float(wave["slope_score"]), 2),
                "drawdown_score": round(float(wave["drawdown_score"]), 2),
                "stability_score": round(float(wave["stability_score"]), 2),
                "volume_score": round(float(wave["volume_score"]), 2),
                "total_score": round(float(wave["total_score"]), 2),
                "historical_percentile": round(float(wave["historical_percentile"]), 2),
            }
        )
    return rows


def _interval_rows(group: str, prices: pd.DataFrame) -> list[dict[str, Any]]:
    unique_dates = prices["trade_date"].sort_values().drop_duplicates()
    if unique_dates.empty:
        return []

    start_date = unique_dates.iloc[max(0, len(unique_dates) - RECENT_WINDOW_DAYS)]
    end_date = unique_dates.iloc[-1]
    ranked = score_interval_continuity(prices, start_date, end_date)
    if ranked.empty:
        return []

    names = {asset.ts_code: asset.name for asset in get_assets(group)}
    rows: list[dict[str, Any]] = []
    for _, row in ranked.iterrows():
        rows.append(
            {
                "group": group,
                "group_name": ASSET_GROUPS[group],
                "symbol": row["symbol"],
                "name": names.get(row["symbol"], row["symbol"]),
                "direction": row["direction"],
                "direction_name": "上涨" if row["direction"] == "up" else "下跌",
                "start_date": pd.Timestamp(row["start_date"]).strftime("%Y-%m-%d"),
                "end_date": pd.Timestamp(row["end_date"]).strftime("%Y-%m-%d"),
                "pct_change": round(float(row["pct_change"]), 2),
                "points": round(float(row["points"]), 2),
                "trend_day_ratio": round(float(row["trend_day_ratio"]), 2),
                "adverse_day_ratio": round(float(row["adverse_day_ratio"]), 2),
                "max_adverse_pct": round(float(row["max_adverse_pct"]), 2),
                "stability_score": round(float(row["stability_score"]), 2),
                "consistency_score": round(float(row["consistency_score"]), 2),
                "interval_score": round(float(row["interval_score"]), 2),
            }
        )
    return rows


def _group_summary(assets: pd.DataFrame) -> pd.DataFrame:
    if assets.empty:
        return pd.DataFrame(
            columns=[
                "group",
                "group_name",
                "symbols",
                "rows",
                "waves",
                "median_reversal",
                "median_reversal_pct",
                "median_reversal_vol_multiple",
                "avg_days",
                "waves_per_year",
                "annual_volatility",
                "total_return",
                "top_score",
            ]
        )
    return (
        assets.groupby(["group", "group_name"], as_index=False)
        .agg(
            symbols=("symbol", "count"),
            rows=("rows", "sum"),
            waves=("waves", "sum"),
            median_reversal=("min_reversal", "median"),
            median_reversal_pct=("min_reversal_pct", "median"),
            median_reversal_vol_multiple=("reversal_vol_multiple", "median"),
            avg_days=("avg_days", "mean"),
            waves_per_year=("waves_per_year", "mean"),
            annual_volatility=("annual_volatility", "median"),
            total_return=("total_return", "median"),
            top_score=("top_score", "max"),
        )
        .round(
            {
                "median_reversal": 2,
                "median_reversal_pct": 2,
                "median_reversal_vol_multiple": 2,
                "avg_days": 1,
                "waves_per_year": 2,
                "annual_volatility": 2,
                "total_return": 2,
                "top_score": 2,
            }
        )
    )


def _score_correlations(waves: pd.DataFrame) -> pd.DataFrame:
    features = [
        ("abs_pct_change", "绝对涨跌幅"),
        ("points", "波段点数"),
        ("days", "持续天数"),
        ("points_per_day", "点数斜率"),
        ("strength_score", "强度分"),
        ("duration_score", "持续分"),
        ("slope_score", "斜率分"),
        ("drawdown_score", "回撤控制分"),
        ("stability_score", "稳定性分"),
        ("volume_score", "量能配合分"),
    ]
    rows: list[dict[str, Any]] = []
    if waves.empty:
        return pd.DataFrame(columns=["feature", "feature_name", "correlation", "abs_correlation"])

    target = pd.to_numeric(waves["total_score"], errors="coerce")
    for feature, feature_name in features:
        values = pd.to_numeric(waves[feature], errors="coerce") if feature in waves else pd.Series(dtype=float)
        valid = pd.concat([values, target], axis=1).dropna()
        if len(valid) < 3 or valid.iloc[:, 0].nunique() < 2 or valid.iloc[:, 1].nunique() < 2:
            corr = 0.0
        else:
            corr = float(valid.iloc[:, 0].corr(valid.iloc[:, 1]))
            if pd.isna(corr):
                corr = 0.0
        rows.append(
            {
                "feature": feature,
                "feature_name": feature_name,
                "correlation": round(corr, 3),
                "abs_correlation": round(abs(corr), 3),
            }
        )
    return pd.DataFrame(rows).sort_values("abs_correlation", ascending=False).reset_index(drop=True)


def _select_typical_waves(waves: pd.DataFrame) -> pd.DataFrame:
    if waves.empty:
        return pd.DataFrame(columns=["case_type", *waves.columns.tolist()])

    candidates: list[pd.Series] = []
    used_keys: set[tuple[str, str, str]] = set()

    def add_case(case_type: str, frame: pd.DataFrame) -> None:
        if frame.empty:
            return
        row = frame.iloc[0].copy()
        key = (str(row["symbol"]), str(row["start_date"]), str(row["end_date"]))
        if key in used_keys:
            return
        row["case_type"] = case_type
        used_keys.add(key)
        candidates.append(row)

    add_case("高分波段", waves.sort_values("total_score", ascending=False))
    add_case("低分反例", waves.sort_values(["total_score", "days"], ascending=[True, False]))
    add_case("快速推进", waves.sort_values(["pct_per_day", "total_score"], ascending=False))
    add_case("长周期顺滑", waves.sort_values(["days", "stability_score", "total_score"], ascending=False))
    add_case("回撤控制样本", waves.sort_values(["drawdown_score", "total_score"], ascending=False))

    return pd.DataFrame(candidates).reset_index(drop=True)


def _typical_wave_paths(price_frames: dict[str, pd.DataFrame], typical_waves: pd.DataFrame) -> pd.DataFrame:
    rows: list[dict[str, Any]] = []
    if typical_waves.empty:
        return pd.DataFrame(columns=["case_type", "symbol", "name", "progress", "trade_date", "relative_close"])

    for _, wave in typical_waves.head(5).iterrows():
        prices = price_frames.get(str(wave["symbol"]))
        if prices is None or prices.empty:
            continue
        segment = prices[
            (prices["trade_date"] >= pd.Timestamp(wave["start_date"]))
            & (prices["trade_date"] <= pd.Timestamp(wave["end_date"]))
        ].sort_values("trade_date")
        if segment.empty:
            continue
        start = float(segment.iloc[0]["close"])
        denominator = max(len(segment) - 1, 1)
        for index, point in segment.reset_index(drop=True).iterrows():
            rows.append(
                {
                    "case_type": wave["case_type"],
                    "symbol": wave["symbol"],
                    "name": wave["name"],
                    "progress": round(index / denominator * 100.0, 2),
                    "trade_date": pd.Timestamp(point["trade_date"]).strftime("%Y-%m-%d"),
                    "relative_close": round(float(point["close"]) / start * 100.0, 2) if start else 100.0,
                }
            )
    return pd.DataFrame(rows)


def render_charts(analysis: dict[str, Any], chart_dir: Path) -> dict[str, Path]:
    _configure_matplotlib()
    paths = {
        "threshold_sensitivity": chart_dir / "threshold_vs_wave_count.png",
        "group_distribution": chart_dir / "group_wave_distribution.png",
        "score_drivers": chart_dir / "score_driver_correlations.png",
        "interval_ranking": chart_dir / "recent_interval_ranking.png",
        "typical_paths": chart_dir / "typical_wave_paths.png",
    }
    _plot_threshold_sensitivity(analysis["assets"], paths["threshold_sensitivity"])
    _plot_group_distribution(analysis["wave_distribution"], paths["group_distribution"])
    _plot_score_drivers(analysis["score_correlations"], paths["score_drivers"])
    _plot_interval_ranking(analysis["intervals"], paths["interval_ranking"])
    _plot_typical_paths(analysis["typical_paths"], paths["typical_paths"])
    return paths


def _configure_matplotlib() -> None:
    plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "Arial Unicode MS", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams["figure.facecolor"] = "white"
    plt.rcParams["axes.facecolor"] = "white"


def _plot_threshold_sensitivity(assets: pd.DataFrame, output_path: Path) -> None:
    fig, ax = plt.subplots(figsize=(8.2, 5.0), dpi=170)
    if assets.empty:
        _empty_chart(ax)
    else:
        colors = {"指数": "#2563eb", "板块": "#b7791f", "商品": "#0f766e"}
        for group_name, frame in assets.groupby("group_name"):
            ax.scatter(
                frame["min_reversal_pct"],
                frame["waves"],
                s=np.clip(frame["annual_volatility"].fillna(0) * 2.5, 45, 220),
                alpha=0.78,
                color=colors.get(group_name, "#667085"),
                label=group_name,
                edgecolor="white",
                linewidth=0.7,
            )
            for _, row in frame.iterrows():
                ax.annotate(str(row["name"]), (row["min_reversal_pct"], row["waves"]), xytext=(4, 4), textcoords="offset points", fontsize=7)
        ax.set_xlabel("自适应反转阈值（折算为收盘价百分比）")
        ax.set_ylabel("历史波段数量")
        ax.legend(frameon=False, fontsize=8)
    ax.set_title("阈值敏感性：百分比阈值与波段数量", fontsize=13, pad=12)
    _polish_axes(ax)
    fig.tight_layout()
    fig.savefig(output_path, bbox_inches="tight")
    plt.close(fig)


def _plot_group_distribution(waves: pd.DataFrame, output_path: Path) -> None:
    fig, axes = plt.subplots(1, 2, figsize=(9.2, 4.8), dpi=170)
    metrics = [("days", "波段持续天数"), ("total_score", "趋势评分")]
    if waves.empty:
        for ax in axes:
            _empty_chart(ax)
    else:
        groups = [name for name in ASSET_GROUPS.values() if name in set(waves["group_name"])]
        for ax, (metric, title) in zip(axes, metrics):
            data = [waves.loc[waves["group_name"] == group, metric].dropna().to_numpy() for group in groups]
            ax.boxplot(data, tick_labels=groups, patch_artist=True, boxprops={"facecolor": "#dbeafe", "edgecolor": "#2563eb"}, medianprops={"color": "#b7791f"})
            ax.set_title(title, fontsize=11)
            _polish_axes(ax)
    fig.suptitle("资产组波段分布：持续时间与评分离散度", fontsize=13)
    fig.tight_layout()
    fig.savefig(output_path, bbox_inches="tight")
    plt.close(fig)


def _plot_score_drivers(correlations: pd.DataFrame, output_path: Path) -> None:
    fig, ax = plt.subplots(figsize=(8.4, 4.8), dpi=170)
    if correlations.empty:
        _empty_chart(ax)
    else:
        frame = correlations.sort_values("abs_correlation", ascending=True).tail(8)
        values = frame["correlation"].to_numpy()
        labels = frame["feature_name"].tolist()
        colors = ["#0f766e" if value >= 0 else "#b91c1c" for value in values]
        ax.barh(labels, values, color=colors, alpha=0.86)
        ax.axvline(0, color="#98a2b3", linewidth=1)
        for index, value in enumerate(values):
            ax.text(value, index, f" {value:.2f}", va="center", ha="left" if value >= 0 else "right", fontsize=8)
        ax.set_xlim(-1, 1)
        ax.set_xlabel("与总评分的相关系数")
    ax.set_title("评分驱动相关性：哪些维度真正推高总分", fontsize=13, pad=12)
    _polish_axes(ax)
    fig.tight_layout()
    fig.savefig(output_path, bbox_inches="tight")
    plt.close(fig)


def _plot_interval_ranking(intervals: pd.DataFrame, output_path: Path) -> None:
    fig, ax = plt.subplots(figsize=(8.8, 5.2), dpi=170)
    if intervals.empty:
        _empty_chart(ax)
    else:
        top = intervals.head(5)
        bottom = intervals.tail(5)
        frame = pd.concat([top, bottom]).drop_duplicates(subset=["symbol"]).sort_values("interval_score")
        labels = (frame["group_name"] + " / " + frame["name"]).tolist()
        colors = ["#2563eb" if value >= frame["interval_score"].median() else "#98a2b3" for value in frame["interval_score"]]
        ax.barh(labels, frame["interval_score"], color=colors, alpha=0.88)
        for index, (_, row) in enumerate(frame.iterrows()):
            ax.text(
                row["interval_score"],
                index,
                f" {row['interval_score']:.1f} | {row['direction_name']} {row['pct_change']:.1f}%",
                va="center",
                ha="left",
                fontsize=8,
            )
        ax.set_xlim(0, 105)
        ax.set_xlabel("区间连续性评分")
    ax.set_title(f"最近约 {RECENT_WINDOW_DAYS} 个交易日连续性排名：Top 与 Bottom", fontsize=13, pad=12)
    _polish_axes(ax)
    fig.tight_layout()
    fig.savefig(output_path, bbox_inches="tight")
    plt.close(fig)


def _plot_typical_paths(paths: pd.DataFrame, output_path: Path) -> None:
    fig, ax = plt.subplots(figsize=(8.8, 5.0), dpi=170)
    if paths.empty:
        _empty_chart(ax)
    else:
        palette = ["#2563eb", "#b7791f", "#0f766e", "#7c3aed", "#b91c1c"]
        for color, (case_type, frame) in zip(palette, paths.groupby("case_type", sort=False)):
            label = f"{case_type} / {frame.iloc[0]['name']}"
            ax.plot(frame["progress"], frame["relative_close"], linewidth=2.0, label=label, color=color)
        ax.axhline(100, color="#98a2b3", linewidth=1, linestyle="--")
        ax.set_xlabel("波段进程百分比")
        ax.set_ylabel("相对起点收盘价 = 100")
        ax.legend(frameon=False, fontsize=8, loc="best")
    ax.set_title("典型波段路径对比：高分、低分与不同质量来源", fontsize=13, pad=12)
    _polish_axes(ax)
    fig.tight_layout()
    fig.savefig(output_path, bbox_inches="tight")
    plt.close(fig)


def _empty_chart(ax: plt.Axes) -> None:
    ax.text(0.5, 0.5, "暂无可用数据", ha="center", va="center", transform=ax.transAxes, color="#667085")
    ax.set_xticks([])
    ax.set_yticks([])


def _polish_axes(ax: plt.Axes) -> None:
    ax.grid(axis="x", color="#e5e7eb", linewidth=0.8)
    ax.set_axisbelow(True)
    for spine in ["top", "right", "left"]:
        ax.spines[spine].set_visible(False)
    ax.spines["bottom"].set_color("#d0d5dd")
    ax.tick_params(axis="both", labelsize=8, colors="#344054")
    ax.xaxis.label.set_color("#475467")
    ax.yaxis.label.set_color("#475467")


def write_docx_report(analysis: dict[str, Any], chart_paths: dict[str, Path], output_path: Path) -> None:
    document = Document()
    _set_default_font(document)

    title = document.add_heading(REPORT_TITLE, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = document.add_paragraph(f"生成时间：{analysis['generated_at']}；数据目录：{analysis['data_dir']}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_executive_summary(document, analysis)
    _add_research_question(document)
    _add_data_and_method(document, analysis)
    _add_wave_identification(document, analysis, chart_paths)
    _add_score_driver_analysis(document, analysis, chart_paths)
    _add_group_characteristics(document, analysis, chart_paths)
    _add_current_interval_observation(document, analysis, chart_paths)
    _add_limits_and_roadmap(document, analysis)
    _add_appendix(document, analysis)

    document.save(output_path)


def _set_default_font(document: Document) -> None:
    styles = document.styles
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        if style_name == "Normal":
            style.font.size = Pt(10.5)


def _add_executive_summary(document: Document, analysis: dict[str, Any]) -> None:
    assets = analysis["assets"]
    waves = analysis["wave_distribution"]
    correlations = analysis["score_correlations"]
    intervals = analysis["intervals"]
    total_assets = int(len(assets))
    total_rows = int(assets["rows"].sum()) if not assets.empty else 0
    total_waves = int(len(waves))
    top_driver = correlations.iloc[0] if not correlations.empty else None
    interval_leader = intervals.iloc[0] if not intervals.empty else None

    document.add_heading("Executive Summary", level=1)
    _add_bullet(
        document,
        f"本报告复盘本地 CSV 覆盖的 {total_assets} 个品种、{total_rows:,} 行日线数据，识别 {total_waves:,} 个历史波段。系统的定位不是预测下一段行情，而是把历史趋势路径拆成可比较、可复盘的质量样本。",
    )
    if not assets.empty:
        median_reversal = float(assets["min_reversal"].median())
        median_reversal_pct = float(assets["min_reversal_pct"].median())
        median_vol_multiple = float(assets["reversal_vol_multiple"].median())
        median_waves_per_year = float(assets["waves_per_year"].median())
        _add_bullet(
            document,
            f"当前自适应反转阈值的中位数为 {median_reversal:.2f} 点，折算为价格比例约 {median_reversal_pct:.2f}%，约等于 {median_vol_multiple:.1f} 倍中位日波动；单品种年化识别波段数中位数约 {median_waves_per_year:.1f} 段。",
        )
    if top_driver is not None:
        _add_bullet(
            document,
            f"评分驱动上，当前样本与总评分相关度最高的是“{top_driver['feature_name']}”（相关系数 {top_driver['correlation']:.2f}）。这说明总分不是单纯涨跌幅排名，而是在路径质量维度上重新分层。",
        )
    if interval_leader is not None:
        _add_bullet(
            document,
            f"最近约 {RECENT_WINDOW_DAYS} 个交易日的横向比较中，{interval_leader['group_name']} / {interval_leader['name']} 暂列区间连续性第一，方向为{interval_leader['direction_name']}，区间评分 {interval_leader['interval_score']:.1f}。",
        )


def _add_research_question(document: Document) -> None:
    document.add_heading("1. 研究背景与问题定义", level=1)
    document.add_paragraph(
        "连续性系统要解决的问题，是把“看起来在涨或在跌”的价格序列，转化成一组可以横向比较的历史波段。它关注的不是某一天的信号，而是趋势从起点到终点的推进质量：是否走得够远、是否持续、是否顺滑、反向波动是否可控、量能是否配合。"
    )
    document.add_paragraph(
        "因此，本报告把系统定义为“历史趋势路径质量的分层评分器”。它可以服务于复盘、样本筛选和品种横向观察，但不直接给出买卖建议，也不把高分等同于未来必然延续。"
    )


def _add_data_and_method(document: Document, analysis: dict[str, Any]) -> None:
    assets = analysis["assets"]
    group_summary = analysis["group_summary"]
    document.add_heading("2. 数据样本与方法口径", level=1)
    if assets.empty:
        document.add_paragraph("当前数据目录下没有可用 CSV，报告仅保留方法结构。")
        return

    start_date = assets["start_date"].min()
    end_date = assets["end_date"].max()
    groups = "、".join(group_summary["group_name"].tolist()) if not group_summary.empty else "指数、板块、商品"
    document.add_paragraph(
        f"样本来自本地数据目录，覆盖 {groups} 三类资产，最早日期 {start_date}，最新日期 {end_date}。各品种统一使用 OHLCV 字段，先在全历史上识别并评分，再按展示日期筛选，避免历史分位随网页筛选窗口漂移。"
    )
    document.add_paragraph(
        "反转阈值支持三种口径：自动模式按 ATR、日收益波动和绝对日波动估算；百分比模式按价格比例判断反转，更适合跨资产理解；点数模式保留给熟悉单一品种的人工精调。报告中同时展示点数、折算百分比和相对日波动倍数，避免不同价格量级的品种被点数误导。"
    )


def _add_wave_identification(document: Document, analysis: dict[str, Any], chart_paths: dict[str, Path]) -> None:
    assets = analysis["assets"]
    document.add_heading("3. 波段识别是否合理", level=1)
    _add_picture(document, chart_paths["threshold_sensitivity"])
    if assets.empty:
        document.add_paragraph("由于没有可用资产统计，无法评价阈值敏感性。")
        return

    corr = _safe_corr(assets["min_reversal_pct"], assets["waves"])
    high_freq = assets.sort_values("waves_per_year", ascending=False).head(1).iloc[0]
    smooth = assets.sort_values("waves_per_year", ascending=True).head(1).iloc[0]
    document.add_paragraph(
        f"折算百分比阈值与波段数量的相关系数为 {corr:.2f}。如果该值显著为负，说明阈值越高切分越少；如果接近 0，则更多由品种自身波动形态决定。当前年化波段最密集的是 {high_freq['name']}（{high_freq['waves_per_year']:.1f} 段/年），最平滑的是 {smooth['name']}（{smooth['waves_per_year']:.1f} 段/年）。"
    )
    document.add_paragraph(
        "这个结果用于判断识别是否过度切分：如果某品种波段数量高、平均持续天数短、但总收益和波动并不高，就需要提高百分比阈值或增加最小持续天数；反过来，如果长期只有少量波段，系统可能过度平滑，容易漏掉中级别趋势。"
    )


def _add_score_driver_analysis(document: Document, analysis: dict[str, Any], chart_paths: dict[str, Path]) -> None:
    correlations = analysis["score_correlations"]
    typical = analysis["typical_waves"]
    document.add_heading("4. 趋势评分的真实驱动因素", level=1)
    _add_picture(document, chart_paths["score_drivers"])
    if correlations.empty:
        document.add_paragraph("当前波段样本不足，无法稳定计算评分驱动相关性。")
    else:
        top = correlations.head(3)
        driver_text = "、".join(f"{row.feature_name}（{row.correlation:.2f}）" for row in top.itertuples())
        document.add_paragraph(
            f"相关性分析显示，当前评分最主要的解释变量是 {driver_text}。这里的重点不是把相关系数当作因果，而是检查评分体系有没有被单一涨跌幅绑架。只要回撤控制、稳定性、斜率和持续维度能进入前列，系统就更接近“趋势路径质量”而不是“涨跌幅榜”。"
        )
        abs_return_corr = correlations.loc[correlations["feature"] == "abs_pct_change", "correlation"]
        if not abs_return_corr.empty:
            document.add_paragraph(
                f"绝对涨跌幅与总评分的相关系数为 {float(abs_return_corr.iloc[0]):.2f}。这项指标应该重要，但不应独占解释力；否则低回撤、慢推进但质量很高的波段会被系统低估。"
            )
    if not typical.empty:
        document.add_paragraph("下表只保留典型样本，用来解释不同质量来源，而不是复述所有 Top 排名。")
        _add_dataframe_table(
            document,
            typical,
            ["case_type", "group_name", "symbol", "name", "direction_name", "level", "days", "pct_change", "drawdown_score", "stability_score", "total_score"],
            max_rows=5,
        )
        _add_picture(document, chart_paths["typical_paths"])
        document.add_paragraph(
            "典型路径图把不同长度的波段归一到 0-100% 的进程，重点看推进过程，而不是只看终点涨跌幅。高分样本通常不只是终点更远，还会表现出更少的中途折返或更清晰的斜率。"
        )


def _add_group_characteristics(document: Document, analysis: dict[str, Any], chart_paths: dict[str, Path]) -> None:
    group_summary = analysis["group_summary"]
    waves = analysis["wave_distribution"]
    document.add_heading("5. 不同资产组的趋势特征", level=1)
    _add_picture(document, chart_paths["group_distribution"])
    if group_summary.empty:
        document.add_paragraph("资产组统计为空，无法做结构差异比较。")
        return

    for _, row in group_summary.iterrows():
        group_waves = waves[waves["group_name"] == row["group_name"]] if not waves.empty else pd.DataFrame()
        if group_waves.empty:
            document.add_paragraph(f"{row['group_name']}：当前没有识别到足够波段。")
            continue
        up_ratio = float((group_waves["direction"] == "up").mean() * 100)
        median_score = float(group_waves["total_score"].median())
        document.add_paragraph(
            f"{row['group_name']}：覆盖 {int(row['symbols'])} 个品种，合计 {int(row['waves'])} 个波段，中位总收益 {row['total_return']:.1f}%，中位年化波动 {row['annual_volatility']:.1f}%。上涨波段占比 {up_ratio:.1f}%，波段评分中位数 {median_score:.1f}。"
        )
    document.add_paragraph(
        "宽基指数通常更适合观察系统性趋势的连续性，板块与商品更容易出现强波段和尖锐反转。因此，跨资产比较时不能只看绝对点数，应更多依赖同品种分位、方向一致性和波动率归一化后的结果。"
    )


def _add_current_interval_observation(document: Document, analysis: dict[str, Any], chart_paths: dict[str, Path]) -> None:
    intervals = analysis["intervals"]
    document.add_heading("6. 当前区间连续性观察", level=1)
    _add_picture(document, chart_paths["interval_ranking"])
    if intervals.empty:
        document.add_paragraph("当前数据无法形成最近区间排名。")
        return

    top = intervals.head(3)
    bottom = intervals.tail(3)
    top_names = "、".join(f"{row.name}（{row.interval_score:.1f}）" for row in top.itertuples())
    bottom_names = "、".join(f"{row.name}（{row.interval_score:.1f}）" for row in bottom.itertuples())
    document.add_paragraph(
        f"最近约 {RECENT_WINDOW_DAYS} 个交易日中，连续性靠前的是 {top_names}；靠后的是 {bottom_names}。这里的排名综合了涨跌幅、斜率、回撤、稳定性、量能和方向一致性，因此比单纯区间涨跌幅更能反映推进过程是否顺畅。"
    )
    document.add_paragraph(
        "方向一致性是本轮优化后新增的关键证据：同样上涨 10%，若其中大量交易日反向运行，只靠尾端拉升完成涨幅，连续性评分就应低于稳步推进的品种。"
    )
    _add_dataframe_table(
        document,
        pd.concat([top, bottom]).drop_duplicates(subset=["symbol"]),
        ["group_name", "symbol", "name", "direction_name", "pct_change", "trend_day_ratio", "max_adverse_pct", "interval_score"],
        max_rows=6,
    )


def _add_limits_and_roadmap(document: Document, analysis: dict[str, Any]) -> None:
    document.add_heading("7. 模型局限与优化路线", level=1)
    _add_numbered(document, "阈值仍是系统最敏感的入口。本轮新增百分比口径后，下一步可以继续加入 ATR 倍数或历史波动分位，让指数、板块、商品之间更可比。")
    _add_numbered(document, "量能分数已经加入常量数据保护，但期货连续合约、指数和板块的成交量含义不同，后续应按资产组分别定义量能解释。")
    _add_numbered(document, "当前连续性用方向一致性和回撤度量近似表达，仍可加入连续阳线/阴线、回撤次数、反弹失败次数等更接近交易语言的结构指标。")
    _add_numbered(document, "历史分位已经稳定到全历史口径，但当前波段若尚未结束，仍存在“未完成样本”问题，适合在网页中增加当前波段置信度或未完成标记。")
    _add_numbered(document, "报告结果依赖本地 CSV 的更新频率。若线上网页已接入 GitHub 与 Streamlit Cloud，代码推送后会更新逻辑，但数据仍取决于仓库或云端可访问的数据文件。")


def _add_appendix(document: Document, analysis: dict[str, Any]) -> None:
    document.add_heading("附录", level=1)
    document.add_paragraph("附录保留完整统计口径，正文只引用关键结论，避免图表和表格重复传达同一信息。")

    document.add_heading("附录 A：品种统计", level=2)
    _add_dataframe_table(
        document,
        analysis["assets"],
        [
            "group_name",
            "symbol",
            "name",
            "start_date",
            "end_date",
            "rows",
            "total_return",
            "annual_volatility",
            "min_reversal",
            "min_reversal_pct",
            "reversal_vol_multiple",
            "waves",
            "waves_per_year",
            "avg_days",
            "top_score",
        ],
        max_rows=40,
    )

    document.add_heading("附录 B：高分历史波段", level=2)
    _add_dataframe_table(
        document,
        analysis["top_waves"],
        ["group_name", "symbol", "name", "direction_name", "level", "start_date", "end_date", "days", "pct_change", "total_score", "historical_percentile"],
        max_rows=15,
    )

    document.add_heading("附录 C：典型波段路径", level=2)
    document.add_paragraph("典型路径图已放入正文第 4 节，附录不再重复插图。")


def _add_picture(document: Document, path: Path | None) -> None:
    if path and Path(path).exists():
        document.add_picture(str(path), width=Inches(6.45))


def _add_dataframe_table(document: Document, frame: pd.DataFrame, columns: list[str], max_rows: int | None = None) -> None:
    if frame.empty:
        document.add_paragraph("暂无可用表格数据。")
        return
    available = [column for column in columns if column in frame.columns]
    display = frame.loc[:, available].head(max_rows).copy() if max_rows else frame.loc[:, available].copy()
    table = document.add_table(rows=1, cols=len(available))
    table.style = "Table Grid"
    for index, column in enumerate(available):
        table.rows[0].cells[index].text = _column_label(column)
    for _, row in display.iterrows():
        cells = table.add_row().cells
        for index, column in enumerate(available):
            cells[index].text = _format_cell(row[column])


def _add_bullet(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Bullet")


def _add_numbered(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Number")


def _safe_corr(left: pd.Series, right: pd.Series) -> float:
    frame = pd.concat([pd.to_numeric(left, errors="coerce"), pd.to_numeric(right, errors="coerce")], axis=1).dropna()
    if len(frame) < 3 or frame.iloc[:, 0].nunique() < 2 or frame.iloc[:, 1].nunique() < 2:
        return 0.0
    corr = frame.iloc[:, 0].corr(frame.iloc[:, 1])
    return round(float(corr), 2) if pd.notna(corr) else 0.0


def _column_label(column: str) -> str:
    labels = {
        "case_type": "样本类型",
        "group_name": "资产组",
        "symbols": "品种数",
        "rows": "数据行数",
        "waves": "波段数",
        "median_reversal": "中位反转点数",
        "median_reversal_pct": "中位反转比例(%)",
        "median_reversal_vol_multiple": "中位日波动倍数",
        "avg_days": "平均持续天数",
        "waves_per_year": "年化波段数",
        "annual_volatility": "年化波动率(%)",
        "total_return": "总涨跌幅(%)",
        "top_score": "最高分",
        "symbol": "代码",
        "name": "名称",
        "start_date": "开始日期",
        "end_date": "结束日期",
        "min_reversal": "最小反转点数",
        "min_reversal_pct": "反转比例(%)",
        "reversal_vol_multiple": "日波动倍数",
        "median_points": "中位波段点数",
        "max_points": "最大波段点数",
        "direction_name": "方向",
        "level": "级别",
        "points": "涨跌点数",
        "pct_change": "涨跌幅(%)",
        "days": "天数",
        "drawdown_score": "回撤控制分",
        "stability_score": "稳定性分",
        "total_score": "趋势评分",
        "historical_percentile": "历史分位",
        "trend_day_ratio": "方向一致性",
        "max_adverse_pct": "最大反向波动(%)",
        "interval_score": "区间评分",
    }
    return labels.get(column, column)


def _format_cell(value: Any) -> str:
    if pd.isna(value):
        return ""
    if isinstance(value, float):
        return f"{value:.2f}"
    return str(value)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate the continuity wave research Word report.")
    parser.add_argument("--data-dir", default=str(ROOT_DIR / "data"))
    parser.add_argument("--output", default=str(DEFAULT_OUTPUT))
    parser.add_argument("--chart-dir", default=str(DEFAULT_CHART_DIR))
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    output = generate_report(args.data_dir, args.output, args.chart_dir)
    print(output)


if __name__ == "__main__":
    main()
